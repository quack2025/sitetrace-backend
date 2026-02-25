"""Document parsing pipeline — Word (.docx) and Excel (.xlsx) extraction.

Extracts structured text preserving hierarchy for AI context.
"""
import io
from loguru import logger


MAX_TEXT_LENGTH = 10_000


async def parse_docx(file_bytes: bytes) -> str:
    """Extract structured text from a Word document.

    Preserves headings, paragraphs, and table structure.
    Returns plain text with markdown-like formatting for AI consumption.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para and para.text.strip():
                style_name = para.style.name if para.style else ""
                if "Heading 1" in style_name:
                    parts.append(f"\n# {para.text}\n")
                elif "Heading 2" in style_name:
                    parts.append(f"\n## {para.text}\n")
                elif "Heading 3" in style_name:
                    parts.append(f"\n### {para.text}\n")
                elif "List" in style_name or "Bullet" in style_name:
                    parts.append(f"- {para.text}")
                else:
                    parts.append(para.text)

        elif tag == "tbl":
            # Table
            for table in doc.tables:
                if table._element is element:
                    parts.append(_extract_table(table))
                    break

    text = "\n".join(parts).strip()

    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + "\n\n[... truncated, document too long ...]"
        logger.warning(
            f"DOCX text truncated from {len(text)} to {MAX_TEXT_LENGTH} chars"
        )

    logger.info(f"DOCX parsed: {len(text)} chars, {len(doc.paragraphs)} paragraphs")
    return text


def _extract_table(table) -> str:
    """Extract a Word table as markdown-style text."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))

    if not rows:
        return ""

    # Add header separator after first row
    result = [rows[0]]
    if len(rows) > 1:
        result.append(" | ".join(["---"] * len(table.columns)))
        result.extend(rows[1:])

    return "\n".join(result)


async def parse_xlsx(file_bytes: bytes) -> str:
    """Extract tabular text from an Excel file.

    Extracts all sheets as markdown-style tables.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_parts = [f"\n## Sheet: {sheet_name}\n"]
        row_count = 0

        for row in ws.iter_rows(values_only=True):
            values = [str(cell) if cell is not None else "" for cell in row]
            # Skip completely empty rows
            if not any(v.strip() for v in values):
                continue
            sheet_parts.append(" | ".join(values))
            row_count += 1

            # Add header separator after first data row
            if row_count == 1:
                sheet_parts.append(" | ".join(["---"] * len(values)))

        if row_count > 0:
            parts.extend(sheet_parts)

    wb.close()

    text = "\n".join(parts).strip()

    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + "\n\n[... truncated, spreadsheet too large ...]"
        logger.warning(
            f"XLSX text truncated from {len(text)} to {MAX_TEXT_LENGTH} chars"
        )

    logger.info(
        f"XLSX parsed: {len(text)} chars, {len(wb.sheetnames)} sheets"
    )
    return text
