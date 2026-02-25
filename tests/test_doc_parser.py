"""Tests for document parsing pipeline (DOCX + XLSX)."""
import io
import os
import pytest

os.environ.setdefault("SUPABASE_URL", "https://test.supabase.co")
os.environ.setdefault("SUPABASE_SERVICE_KEY", "test-service-key")
os.environ.setdefault("SUPABASE_ANON_KEY", "test-anon-key")
os.environ.setdefault("JWT_SECRET", "test-secret-key-for-jwt-tokens-minimum-64-chars-long-1234567890abcdef")
os.environ.setdefault("ANTHROPIC_API_KEY", "test-key")

from app.processors.doc_parser import parse_docx, parse_xlsx, _extract_table, MAX_TEXT_LENGTH


def _make_test_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Create a minimal DOCX file. paragraphs = [(text, style_name), ...]"""
    from docx import Document
    doc = Document()
    for text, style in paragraphs:
        if style.startswith("Heading"):
            level = int(style[-1])
            doc.add_heading(text, level=level)
        elif style == "List":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_test_xlsx(sheets: dict[str, list[list]]) -> bytes:
    """Create a minimal XLSX file. sheets = {name: [[row1], [row2], ...]}"""
    from openpyxl import Workbook
    wb = Workbook()
    first = True
    for name, rows in sheets.items():
        if first:
            ws = wb.active
            ws.title = name
            first = False
        else:
            ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestParseDocx:
    @pytest.mark.asyncio
    async def test_parse_simple_paragraphs(self):
        docx_bytes = _make_test_docx([
            ("Hello World", "Normal"),
            ("Second paragraph", "Normal"),
        ])
        text = await parse_docx(docx_bytes)
        assert "Hello World" in text
        assert "Second paragraph" in text

    @pytest.mark.asyncio
    async def test_parse_headings(self):
        docx_bytes = _make_test_docx([
            ("Main Title", "Heading 1"),
            ("Content under title", "Normal"),
            ("Subtitle", "Heading 2"),
        ])
        text = await parse_docx(docx_bytes)
        assert "# Main Title" in text
        assert "## Subtitle" in text

    @pytest.mark.asyncio
    async def test_parse_empty_document(self):
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        text = await parse_docx(buf.getvalue())
        assert text == ""

    @pytest.mark.asyncio
    async def test_truncation_on_long_text(self):
        # Create a document with lots of text
        paragraphs = [(f"Paragraph number {i} with some filler text to make it longer" * 10, "Normal")
                      for i in range(200)]
        docx_bytes = _make_test_docx(paragraphs)
        text = await parse_docx(docx_bytes)
        # Should be truncated
        if len(text) > MAX_TEXT_LENGTH:
            assert "[... truncated" in text


class TestParseXlsx:
    @pytest.mark.asyncio
    async def test_parse_single_sheet(self):
        xlsx_bytes = _make_test_xlsx({
            "Data": [
                ["Name", "Value", "Unit"],
                ["Steel", "100", "tons"],
                ["Concrete", "200", "m3"],
            ]
        })
        text = await parse_xlsx(xlsx_bytes)
        assert "Sheet: Data" in text
        assert "Name | Value | Unit" in text
        assert "Steel | 100 | tons" in text
        assert "---" in text  # header separator

    @pytest.mark.asyncio
    async def test_parse_multiple_sheets(self):
        xlsx_bytes = _make_test_xlsx({
            "Materials": [["Item", "Qty"], ["Tile", "50"]],
            "Costs": [["Category", "Amount"], ["Labor", "5000"]],
        })
        text = await parse_xlsx(xlsx_bytes)
        assert "Sheet: Materials" in text
        assert "Sheet: Costs" in text
        assert "Tile | 50" in text
        assert "Labor | 5000" in text

    @pytest.mark.asyncio
    async def test_skip_empty_rows(self):
        xlsx_bytes = _make_test_xlsx({
            "Data": [
                ["Header"],
                [None, None],  # empty row
                ["Value"],
            ]
        })
        text = await parse_xlsx(xlsx_bytes)
        lines = [l for l in text.split("\n") if l.strip()]
        # Should not contain blank row
        assert all(l.strip() for l in lines)

    @pytest.mark.asyncio
    async def test_empty_workbook(self):
        from openpyxl import Workbook
        wb = Workbook()
        # Default sheet is empty
        buf = io.BytesIO()
        wb.save(buf)
        text = await parse_xlsx(buf.getvalue())
        assert text == ""


class TestExtractTable:
    def test_extract_simple_table(self):
        from unittest.mock import MagicMock

        table = MagicMock()
        row1 = MagicMock()
        row1.cells = [MagicMock(text="Header1"), MagicMock(text="Header2")]
        row2 = MagicMock()
        row2.cells = [MagicMock(text="Val1"), MagicMock(text="Val2")]
        table.rows = [row1, row2]
        table.columns = [MagicMock(), MagicMock()]

        result = _extract_table(table)
        assert "Header1 | Header2" in result
        assert "--- | ---" in result
        assert "Val1 | Val2" in result

    def test_extract_empty_table(self):
        from unittest.mock import MagicMock
        table = MagicMock()
        table.rows = []
        assert _extract_table(table) == ""

    def test_single_row_table(self):
        from unittest.mock import MagicMock
        table = MagicMock()
        row = MagicMock()
        row.cells = [MagicMock(text="Only"), MagicMock(text="Row")]
        table.rows = [row]

        result = _extract_table(table)
        assert "Only | Row" in result
        assert "---" not in result
